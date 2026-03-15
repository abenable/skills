from __future__ import annotations

from datetime import datetime
from pathlib import Path
from zipfile import BadZipFile

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from openpyxl import load_workbook

from time_window import in_time_window, parse_time_window


def resolve_writable_report_path(preferred_path: Path) -> Path:
    try:
        preferred_path.parent.mkdir(parents=True, exist_ok=True)
        if not preferred_path.exists():
            return preferred_path
        with preferred_path.open("ab"):
            pass
        return preferred_path
    except PermissionError:
        timestamp = datetime.now().strftime("%H%M%S")
        return preferred_path.with_name(f"{preferred_path.stem}_{timestamp}{preferred_path.suffix}")


def build_time_window_text(paths: list[Path]) -> str:
    if not paths:
        today = datetime.now().date().isoformat()
        return f"统计时间窗口：{today}"

    dates = sorted(path.stem.replace("international_", "") for path in paths)
    if len(dates) == 1:
        return f"统计时间窗口：{dates[0]}"
    return f"统计时间窗口：{dates[0]} 至 {dates[-1]}"


def load_latest_domestic_rows(
    report_dir: Path,
    report_date: str,
    time_window: str = "",
) -> list[dict[str, str]]:
    candidates = sorted(
        list(report_dir.glob("company_mentions.xlsx")) + list(report_dir.glob(f"company_mentions_{report_date}*.xlsx")),
        key=lambda path: path.stat().st_mtime,
        reverse=True,
    )
    return filter_rows(load_rows_from_candidates(candidates), report_date, time_window)


def load_latest_international_rows(
    report_dir: Path,
    report_date: str,
    time_window: str = "",
) -> list[dict[str, str]]:
    candidates = sorted(
        list(report_dir.glob("international_company_mentions.xlsx"))
        + list(report_dir.glob(f"international_company_mentions_{report_date}*.xlsx")),
        key=lambda path: path.stat().st_mtime,
        reverse=True,
    )
    return filter_rows(load_rows_from_candidates(candidates), report_date, time_window)


def load_rows_from_candidates(candidates: list[Path]) -> list[dict[str, str]]:
    for candidate in candidates:
        try:
            workbook = load_workbook(candidate, read_only=True)
            worksheet = workbook.active
            headers = [cell.value for cell in next(worksheet.iter_rows(min_row=1, max_row=1))]
            rows: list[dict[str, str]] = []
            for values in worksheet.iter_rows(min_row=2, values_only=True):
                row = {str(headers[index]): str(value or "") for index, value in enumerate(values)}
                rows.append(row)
            return rows
        except (BadZipFile, OSError, StopIteration, KeyError, ValueError):
            continue
    return []


def filter_rows(rows: list[dict[str, str]], report_date: str, time_window: str = "") -> list[dict[str, str]]:
    if time_window:
        start, end = parse_time_window(time_window)
        return [
            row
            for row in rows
            if in_time_window(str(row.get("资讯时间", "") or ""), start, end)
        ]

    filtered: list[dict[str, str]] = []
    for row in rows:
        news_time = str(row.get("资讯时间", "") or "")
        if news_time.startswith(report_date):
            filtered.append(row)
    return filtered


def top_international_rows(rows: list[dict[str, str]], limit: int = 5) -> list[dict[str, str]]:
    def sort_key(row: dict[str, str]) -> tuple[int, str]:
        try:
            score = int(str(row.get("事件影响力", "") or 0))
        except ValueError:
            score = 0
        return (-score, str(row.get("资讯时间", "") or ""))

    return sorted(rows, key=sort_key)[:limit]


def write_daily_brief(
    report_dir: Path,
    paths: list[Path],
    domestic_rows: list[dict[str, str]],
    international_rows: list[dict[str, str]],
    time_window: str = "",
) -> Path:
    report_date = datetime.now().date().isoformat()
    output_path = resolve_writable_report_path(report_dir / f"AI资讯简报_{report_date}.docx")

    document = Document()

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("每日AI产业资讯简报")
    title_run.bold = True
    title_run.font.size = Pt(16)

    time_window_paragraph = document.add_paragraph()
    header_text = build_time_window_text(paths)
    if isinstance(time_window, str) and time_window.strip():
        header_text = f"统计时间窗口：{time_window}"
    time_window_run = time_window_paragraph.add_run(header_text)
    time_window_run.font.size = Pt(11)

    heading_domestic = document.add_paragraph()
    heading_domestic_run = heading_domestic.add_run("一、AI领军企业动态")
    heading_domestic_run.bold = True
    heading_domestic_run.font.size = Pt(12)

    if not domestic_rows:
        empty_paragraph = document.add_paragraph()
        empty_paragraph.add_run("今日无命中企业资讯。").font.size = Pt(11)
    else:
        for index, row in enumerate(domestic_rows, start=1):
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.space_after = Pt(8)

            ai_title = row.get("AI标题") or row.get("资讯标题") or "未命名资讯"
            company_name = row.get("企业名", "")
            ai_summary = row.get("AI摘要") or row.get("资讯内容", "")
            source_name = row.get("资讯来源", "")

            title_run = paragraph.add_run(f"{index}. {ai_title}\n")
            title_run.bold = True
            title_run.font.size = Pt(11)

            company_run = paragraph.add_run(f"关联企业：{company_name}\n")
            company_run.font.size = Pt(11)

            content_text = f"资讯内容：{ai_summary}"
            if source_name:
                content_text += f"（{source_name}）"
            content_run = paragraph.add_run(content_text)
            content_run.font.size = Pt(11)

    heading_international = document.add_paragraph()
    heading_international_run = heading_international.add_run("二、全球AI动态")
    heading_international_run.bold = True
    heading_international_run.font.size = Pt(12)

    selected_international_rows = top_international_rows(international_rows, limit=5)

    if not selected_international_rows:
        empty_paragraph = document.add_paragraph()
        empty_paragraph.add_run("今日无国际AI资讯。").font.size = Pt(11)
    else:
        for index, row in enumerate(selected_international_rows, start=1):
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.space_after = Pt(8)
            summary = row.get("AI摘要") or row.get("资讯内容") or ""
            source_name = row.get("资讯来源", "")
            content_text = f"{index}. {summary}"
            if source_name:
                content_text += f"（{source_name}）"
            run = paragraph.add_run(content_text)
            run.font.size = Pt(11)

    document.save(output_path)
    return output_path
